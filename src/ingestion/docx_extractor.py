from typing import List
from src.models.schemas import PageContent, TableData

def extract_docx(file_path: str) -> tuple[List[PageContent], str, List[TableData]]:
    """
    Extracts text and tables from DOCX files using python-docx.
    Returns (pages, full_text, tables).
    """
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx is required for .docx extraction. Run: pip install python-docx")

    doc = docx.Document(file_path)
    full_text_list: List[str] = []
    tables: List[TableData] = []

    for p in doc.paragraphs:
        if p.text.strip():
            full_text_list.append(p.text.strip())

    for table in doc.tables:
        rows_data: List[List[str]] = []
        headers: List[str] = []
        for row_idx, row in enumerate(table.rows):
            row_cells = [cell.text.strip() for cell in row.cells]
            if row_idx == 0:
                headers = row_cells
            else:
                rows_data.append(row_cells)
        if headers or rows_data:
            tables.append(TableData(page_number=1, headers=headers, rows=rows_data))

    full_text = "\n\n".join(full_text_list)
    # DOCX does not natively expose page breaks easily, so represent as single continuous page
    pages = [PageContent(page_number=1, text=full_text)]
    return pages, full_text, tables
